"""
Text extraction service - OCR for images, parsing for documents
"""
import logging
import base64
import requests
import asyncio
import re
from typing import Tuple
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
import io
from config import settings

logger = logging.getLogger(__name__)


class TextExtractionService:
    """Extracts text from images and documents"""
    
    def __init__(self):
        # Service initialized with NVIDIA API settings automatically via config
        pass
    
    async def extract_text(self, file_path: str, file_type: str) -> Tuple[str, str]:
        """
        Extract text from file based on type
        Returns (extracted_text, comma_separated_keywords)
        """
        try:
            if file_type in settings.ALLOWED_IMAGE_TYPES:
                return await self._extract_from_image(file_path)
            elif file_type == "application/pdf":
                return await self._extract_from_pdf(file_path)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self._extract_from_docx(file_path)
            elif file_type == "text/plain":
                return await self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type for text extraction: {file_type}")
                
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            raise
    
    async def _call_kimi_vision(self, base64_img: str, mime_type: str, prompt: str) -> Tuple[str, str]:
        """Helper to invoke Kimi K2.5 VLM, returns (text, keywords)"""
        def make_request():
            import json
            for attempt in range(3):
                try:
                    headers = {
                      "Authorization": f"Bearer {settings.NVIDIA_API_KEY}",
                      "Accept": "text/event-stream"
                    }
                    payload = {
                      "model": settings.VLM_MODEL,
                      "messages": [{
                          "role": "user",
                          "content": [
                              {"type": "text", "text": prompt},
                              {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{base64_img}"}}
                          ]
                      }],
                      "max_tokens": 16384,
                      "temperature": 1.00,
                      "top_p": 1.00,
                      "chat_template_kwargs": {"thinking": True},
                      "stream": True
                    }
                    # Give 10s connect, 600s read headers limit. Once headers are read, stream holds the connection.
                    resp = requests.post("https://integrate.api.nvidia.com/v1/chat/completions", headers=headers, json=payload, stream=True, timeout=(10, 600))
                    resp.raise_for_status()
                    
                    full_content = ""
                    for line in resp.iter_lines():
                        if line:
                            line = line.decode('utf-8')
                            if line.startswith('data: ') and line != 'data: [DONE]':
                                try:
                                    chunk = json.loads(line[6:])
                                    if "delta" in chunk["choices"][0] and "content" in chunk["choices"][0]["delta"]:
                                        token = chunk["choices"][0]["delta"]["content"]
                                        if token:
                                            full_content += token
                                except:
                                    pass
                                    
                    content = full_content
                    break
                except requests.exceptions.ReadTimeout:
                    if attempt == 2:
                        raise
                    import time
                    time.sleep(2)
            
            if not content:
                # If model refused or returned absolutely nothing
                return "", ""
                
            # Parse <keywords> securely
            keywords = ""
            text_content = content
            match = re.search(r'<keywords>(.*?)</keywords>', content, re.DOTALL | re.IGNORECASE)
            if match:
                keywords = match.group(1).strip()
                text_content = content.replace(match.group(0), "").strip()
                
            return text_content, keywords
        
        return await asyncio.to_thread(make_request)

    async def _call_nemotron_ocr(self, base64_img: str, mime_type: str, prompt: str = "Extract all text flawlessly.") -> str:
        """Invoke Nemotron-OCR-v1 for high fidelity document text extraction"""
        def make_request():
            import json
            for attempt in range(3):
                try:
                    headers = {
                      "Authorization": f"Bearer {settings.NEMOTRON_API_KEY}",
                      "Accept": "text/event-stream"
                    }
                    payload = {
                      "model": settings.NEMOTRON_MODEL,
                      "messages": [{
                          "role": "user",
                          "content": [
                              {"type": "text", "text": prompt},
                              {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{base64_img}"}}
                          ]
                      }],
                      "max_tokens": 4096,
                      "temperature": 0.1,  # Prevent hallucination for strict OCR
                      "stream": True # Force streaming to avoid API load balancer timeout hangs
                    }
                    resp = requests.post("https://integrate.api.nvidia.com/v1/chat/completions", headers=headers, json=payload, stream=True, timeout=(10, 600))
                    resp.raise_for_status()
                    
                    full_content = ""
                    for line in resp.iter_lines():
                        if line:
                            line = line.decode('utf-8')
                            if line.startswith('data: ') and line != 'data: [DONE]':
                                try:
                                    chunk = json.loads(line[6:])
                                    if "delta" in chunk["choices"][0] and "content" in chunk["choices"][0]["delta"]:
                                        token = chunk["choices"][0]["delta"]["content"]
                                        if token:
                                            full_content += token
                                except:
                                    pass
                                    
                    return full_content
                except requests.exceptions.ReadTimeout as e:
                    with open("nemotron_debug.log", "a", encoding="utf-8") as debug_file:
                        debug_file.write(f"\n--- [NEMOTRON ATTEMPT {attempt + 1} TIMEOUT] ---\n")
                    if attempt == 2:
                        raise e
                    import time
                    time.sleep(2)
                except Exception as e:
                    with open("nemotron_debug.log", "a", encoding="utf-8") as debug_file:
                        debug_file.write(f"\n--- [NEMOTRON ATTEMPT {attempt + 1} ERROR] ---\n")
                        debug_file.write(f"Exception Type: {type(e).__name__}\n")
                        debug_file.write(f"Exception String: {str(e)}\n")
                        if hasattr(e, 'response') and e.response:
                            debug_file.write(f"Response Status: {e.response.status_code}\n")
                            try:
                                debug_file.write(f"Response Text: {e.response.text}\n")
                            except:
                                pass
                    if attempt == 2:
                        raise e
                    import time
                    time.sleep(2)
            
            return ""
        
        return await asyncio.to_thread(make_request)

    async def _extract_from_image(self, image_path: str) -> Tuple[str, str]:
        """Extract text and semantic description from image using Kimi K2.5"""
        try:
            logger.info(f"🤖 Kimi Vision analyzing image: {image_path}")
            
            # Use PIL to safely downscale image to prevent API timeout and drop packet sizes
            with Image.open(image_path) as img:
                # Convert to RGB to safely drop Alpha channels
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                
                # Max 1024x1024 bounds
                img.thumbnail((1024, 1024), Image.Resampling.LANCZOS)
                
                buffer = io.BytesIO()
                img.save(buffer, format="JPEG", quality=85)
                base64_img = base64.b64encode(buffer.getvalue()).decode('utf-8')
                
            mime_type = "image/jpeg"
                
            prompt = "Extract all text exactly as it appears in this image. Add a line break, and then provide a highly detailed semantic description of the image's subject, contents, style, and colors for search engine indexing. Finally, provide exactly 5-10 comma-separated keywords representing the core entities, strictly enclosed in <keywords> tags at the very bottom."
            
            vlm_text, keywords = await self._call_kimi_vision(base64_img, mime_type, prompt)
            vlm_text = ' '.join(vlm_text.split())  # Clean whitespace
            
            logger.info(f"✅ Kimi extracted {len(vlm_text)} characters and {len(keywords.split(','))} keywords")
            return vlm_text, keywords
            
        except Exception as e:
            logger.error(f"❌ Kimi Vision failed: {str(e)}")
            raise
    
    async def _extract_from_pdf(self, pdf_path: str) -> Tuple[str, str]:
        """
        Extract text from PDF with OCR fallback for scanned PDFs
        
        Strategy:
        1. Try PyMuPDF text extraction first (fast, works for text-based PDFs)
        2. If minimal text found, fall back to OCR (slower, works for scanned PDFs)
        """
        try:
            logger.info(f"📄 Extracting text from PDF: {pdf_path}")
            
            text_parts: list[str] = []
            doc = fitz.open(pdf_path)
            logger.info(f"   PDF has {len(doc)} pages")
            
            # Step 1: Try normal PyMuPDF text extraction
            for page_num in range(len(doc)):
                page = doc[page_num]
                logger.debug(f"   Processing page {page_num + 1}/{len(doc)}...")
                
                # Ensure get_text returns string
                page_text = page.get_text()
                # Normalize to string (get_text can return str, list, or dict depending on args)
                if isinstance(page_text, str):
                    text_parts.append(page_text)
                else:
                    # Fallback: convert to string
                    text_parts.append(str(page_text))
                
                logger.debug(f"     Page {page_num + 1} extracted {len(text_parts[-1])} chars")
            
            # Combine all pages
            full_text = '\n'.join(text_parts)
            full_text = ' '.join(full_text.split())  # Clean whitespace
            
            logger.info(f"   PyMuPDF extracted {len(full_text)} characters from {len(text_parts)} pages")
            
            # Step 2: Check if we got sufficient text
            if len(full_text.strip()) >= 50:  # Threshold: at least 50 chars
                logger.info(f"✅ Sufficient text extracted via PyMuPDF")
                if len(full_text) > 0:
                    preview = full_text[:100].replace('\n', ' ')
                    logger.info(f"   Preview: '{preview}...'")
                doc.close()
                return full_text, ""
            
            # Step 3: Minimal text - fall back to Nemotron OCR for scanned PDF
            logger.warning(f"⚠️  Minimal text found ({len(full_text)} chars) - PDF may be scanned")
            logger.info(f"🔄 Falling back to Nemotron OCR extraction...")
            
            ocr_text_parts: list[str] = []
            
            keyword_parts: list[str] = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                logger.info(f"   🔍 Nemotron OCR on page {page_num + 1}/{len(doc)}...")
                
                try:
                    # Convert PDF page to image at higher resolution for better OCR
                    zoom = 2.0  # 2x zoom
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert PyMuPDF pixmap to PIL Image and compress to fit payload constraints
                    from PIL import Image
                    import io
                    
                    # Ensure RGB mode safely
                    img = Image.frombytes("RGB" if pix.alpha == 0 else "RGBA", [pix.width, pix.height], pix.samples)
                    
                    # Scale down to fit 1024x1024 bounds maximum
                    max_dim = 1024
                    ratio = min(max_dim/img.width, max_dim/img.height)
                    if ratio < 1.0:
                        new_w = int(img.width * ratio)
                        new_h = int(img.height * ratio)
                        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                        
                    # Save as compressed JPEG to avoid NVIDIA API payload timeouts
                    buffer = io.BytesIO()
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    img.save(buffer, format='JPEG', quality=85)
                    img_data = buffer.getvalue()
                    
                    # Convert to base64
                    base64_img = base64.b64encode(img_data).decode('utf-8')
                    
                    prompt = "Extract all text perfectly as it appears. Then provide a semantic description of the page and 5-10 comma-separated keywords representing core entities strictly enclosed in <keywords> tags at the very bottom."
                    
                    page_text, page_keywords = await self._call_kimi_vision(base64_img, "image/jpeg", prompt)
                    page_text = ' '.join(page_text.split())  # Clean whitespace
                    
                    if not page_text:
                        logger.warning(f"      ⚠️  Kimi Vision returned empty string on page {page_num + 1}")

                    logger.info(f"      ✅ Kimi Vision extracted {len(page_text)} chars from page {page_num + 1}")
                    
                    ocr_text_parts.append(page_text)
                    if page_keywords:
                        keyword_parts.append(page_keywords)
                    
                except Exception as e:
                    import traceback
                    logger.warning(f"      ⚠️  Kimi Vision failed on page {page_num + 1}: {str(e)}\\n{traceback.format_exc()}")
                    ocr_text_parts.append("")
            
            doc.close()
            
            # Identify if all pages catastrophically failed
            if len(ocr_text_parts) > 0 and all(page_text == "" for page_text in ocr_text_parts):
                raise Exception("Vision OCR totally exhausted or timed out across all scanned pages.")
            
            # Combine OCR results
            final_text = ' '.join(ocr_text_parts)
            final_text = ' '.join(final_text.split())
            
            final_keywords = ', '.join([k for k in keyword_parts if k])
            
            logger.info(f"✅ Nemotron PDF extraction complete: {len(final_text)} characters total")
            if len(final_text) > 0:
                preview = final_text[:100].replace('\n', ' ')
                logger.info(f"   Preview: '{preview}...'")
            else:
                logger.warning(f"⚠️  No text extracted (PDF may be blank or unreadable)")
            
            return final_text, final_keywords
            
        except Exception as e:
            logger.error(f"❌ PDF text extraction failed: {str(e)}")
            raise
    
    async def _extract_from_docx(self, docx_path: str) -> Tuple[str, str]:
        """Extract text from DOCX"""
        try:
            logger.info(f"Extracting text from DOCX: {docx_path}")
            
            doc = Document(docx_path)
            
            # Extract all paragraphs
            text_parts: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(str(para.text))
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(str(cell.text))
            
            full_text = ' '.join(text_parts)
            full_text = ' '.join(full_text.split())  # Clean whitespace
            
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text, ""
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {str(e)}")
            raise
    
    async def _extract_from_txt(self, txt_path: str) -> Tuple[str, str]:
        """Extract text from TXT file"""
        try:
            logger.info(f"Reading text file: {txt_path}")
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    text = ' '.join(text.split())  # Clean whitespace
                    logger.info(f"Read {len(text)} characters (encoding: {encoding})")
                    return text, ""
                    
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            text = ' '.join(text.split())
            logger.info(f"Read {len(text)} characters (fallback encoding)")
            return text, ""
            
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            raise
